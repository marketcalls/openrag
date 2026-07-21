import io

import pytest
from docling.exceptions import ConversionError, DocumentLoadError
from docx import Document as DocxBuilder
from PIL import Image, ImageDraw

from openrag.modules.documents import pipeline
from openrag.modules.documents.pipeline import (
    IngestFailure,
    IngestTransientFailure,
    ParseProfile,
    build_pdf_pipeline_options,
    parse_bytes,
    parse_document,
)


def build_docx() -> bytes:
    document = DocxBuilder()
    document.add_heading("Flux Capacitor Manual", level=1)
    document.add_paragraph("The flux capacitor requires 1.21 gigawatts of power.")
    document.add_heading("Billing", level=1)
    document.add_paragraph("Invoice 0231 covers the plutonium delivery.")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def scanned_pdf(text: str, *, resolution: float = 180.0) -> bytes:
    source = Image.new("RGB", (620, 100), "white")
    ImageDraw.Draw(source).text((10, 35), text, fill="black")
    enlarged = source.resize((2480, 400), Image.Resampling.NEAREST)
    buffer = io.BytesIO()
    enlarged.save(buffer, format="PDF", resolution=resolution)
    return buffer.getvalue()


def test_parse_docx_extracts_blocks_with_kinds() -> None:
    blocks = parse_bytes(build_docx(), "manual.docx")
    text = " ".join(block.text for block in blocks)

    assert "1.21 gigawatts" in text
    assert "Invoice 0231" in text
    assert any(block.kind == "heading" for block in blocks)
    assert all(block.page >= 1 for block in blocks)


def test_parse_txt_fast_path() -> None:
    blocks = parse_bytes(b"para one\n\npara two", "notes.txt")

    assert [block.text for block in blocks] == ["para one", "para two"]
    assert all(block.kind == "text" and block.page == 1 for block in blocks)


def test_empty_file_fails_with_reason() -> None:
    with pytest.raises(IngestFailure, match="empty"):
        parse_bytes(b"", "empty.txt")


def test_unsupported_format_fails_with_reason() -> None:
    with pytest.raises(IngestFailure):
        parse_bytes(b"\x00\x01garbage", "weird.xyz")


def test_docling_runtime_failure_remains_retryable(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    class FailingConverter:
        def __init__(self, **kwargs: object) -> None:
            pass

        def convert(self, *args: object, **kwargs: object) -> object:
            raise RuntimeError("temporary parser initialization failure")

    monkeypatch.setattr(
        "docling.document_converter.DocumentConverter",
        FailingConverter,
    )

    with pytest.raises(IngestTransientFailure, match="temporarily unavailable"):
        parse_document(b"valid container bytes", "manual.docx")


def test_docling_document_load_failure_is_terminal(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    class InvalidConverter:
        def __init__(self, **kwargs: object) -> None:
            pass

        def convert(self, *args: object, **kwargs: object) -> object:
            outer = ConversionError("conversion failed")
            outer.__cause__ = DocumentLoadError("invalid input")
            raise outer

    monkeypatch.setattr(
        "docling.document_converter.DocumentConverter",
        InvalidConverter,
    )

    with pytest.raises(IngestFailure, match="unsupported or unparsable"):
        parse_document(b"invalid container bytes", "manual.docx")


def test_pdf_profile_configures_bounded_rapidocr() -> None:
    profile = ParseProfile(
        ocr_mode="force",
        ocr_languages=("english",),
        ocr_min_confidence=0.61,
        ocr_text_score=0.3,
        ocr_bitmap_area_threshold=0.1,
        ocr_batch_size=2,
        timeout_seconds=45,
    )

    options = build_pdf_pipeline_options(profile)

    assert options.do_ocr is True
    assert options.document_timeout == 45
    assert options.generate_parsed_pages is True
    assert options.ocr_batch_size == 2
    assert options.ocr_options.force_full_page_ocr is True
    assert options.ocr_options.lang == ["english"]
    assert options.ocr_options.bitmap_area_threshold == 0.1
    assert options.ocr_options.text_score == 0.3


def test_real_scanned_pdf_is_ocrd_with_page_confidence() -> None:
    parsed = parse_document(
        scanned_pdf("INVOICE 0231 TOTAL 1250 USD"),
        "invoice.pdf",
        ParseProfile(ocr_mode="force", timeout_seconds=60),
    )

    text = " ".join(block.text for block in parsed.blocks)
    assert "INVOICE" in text
    assert "0231" in text
    assert parsed.page_count == 1
    assert parsed.ocr_pages == (1,)
    assert all(block.extraction_method == "ocr" for block in parsed.blocks)
    assert all(
        block.ocr_confidence is not None and block.ocr_confidence >= 0.5
        for block in parsed.blocks
    )


def test_pdf_page_count_is_rejected_before_conversion() -> None:
    first = Image.new("RGB", (20, 20), "white")
    second = Image.new("RGB", (20, 20), "white")
    buffer = io.BytesIO()
    first.save(buffer, format="PDF", save_all=True, append_images=[second])

    with pytest.raises(IngestFailure, match="page limit"):
        parse_document(
            buffer.getvalue(),
            "two-pages.pdf",
            ParseProfile(max_pages=1),
        )


def test_pdf_rendered_pixel_budget_is_rejected_before_conversion() -> None:
    tiny = Image.new("RGB", (20, 20), "white")
    buffer = io.BytesIO()
    tiny.save(buffer, format="PDF", resolution=1.0)

    with pytest.raises(IngestFailure, match="pixel limit"):
        parse_document(
            buffer.getvalue(),
            "huge-media-box.pdf",
            ParseProfile(max_page_pixels=1_000_000, render_dpi=200),
        )


def test_large_text_native_pdf_uses_fast_path_and_keeps_every_page(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    class FakeTextPage:
        def get_text_range(self) -> str:
            return "Native searchable text " * 20

        def close(self) -> None:
            return None

    class FakePage:
        def get_size(self) -> tuple[int, int]:
            return (612, 792)

        def get_textpage(self) -> FakeTextPage:
            return FakeTextPage()

        def close(self) -> None:
            return None

    class FakeDocument:
        def __len__(self) -> int:
            return 20

        def __getitem__(self, _index: int) -> FakePage:
            return FakePage()

        def __iter__(self):
            return iter(FakePage() for _ in range(20))

        def close(self) -> None:
            return None

    monkeypatch.setattr(pipeline.pdfium, "PdfDocument", lambda _data: FakeDocument())

    parsed = parse_document(b"%PDF-fake", "long-guide.pdf")

    assert parsed.page_count == 20
    assert [block.page for block in parsed.blocks] == list(range(1, 21))
    assert parsed.ocr_pages == ()
    assert all(block.extraction_method == "parser" for block in parsed.blocks)


def test_medium_native_pdf_fast_path_allows_blank_vector_back_cover(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    class FakeTextPage:
        def __init__(self, text: str) -> None:
            self._text = text

        def get_text_range(self) -> str:
            return self._text

        def close(self) -> None:
            return None

    class FakePage:
        def __init__(self, text: str) -> None:
            self._text = text

        def get_size(self) -> tuple[int, int]:
            return (612, 792)

        def get_textpage(self) -> FakeTextPage:
            return FakeTextPage(self._text)

        def get_objects(self, *, filter: list[int]) -> list[object]:
            del filter
            return []

        def close(self) -> None:
            return None

    class FakeDocument:
        pages = ["Native searchable text " * 20 for _ in range(12)] + [""]

        def __len__(self) -> int:
            return len(self.pages)

        def __getitem__(self, index: int) -> FakePage:
            return FakePage(self.pages[index])

        def __iter__(self):
            return iter(FakePage(text) for text in self.pages)

        def close(self) -> None:
            return None

    monkeypatch.setattr(pipeline.pdfium, "PdfDocument", lambda _data: FakeDocument())

    parsed = parse_document(b"%PDF-fake", "team-guide.pdf")

    assert parsed.page_count == 13
    assert [block.page for block in parsed.blocks] == list(range(1, 13))
    assert parsed.ocr_pages == ()
