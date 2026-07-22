import io

import pytest
from docx import Document as DocxBuilder
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from openrag.modules.documents.ingest import (
    run_chunk,
    run_embed_upsert,
    run_parse,
)
from openrag.modules.documents.models import IngestJob
from openrag.modules.documents.pipeline import IngestFailure
from openrag.modules.documents.service import create_from_upload
from openrag.modules.retrieval.service import retrieve
from tests.modules.retrieval.test_retrieve import seed_workspace


def fixture_docx() -> bytes:
    document = DocxBuilder()
    document.add_heading("Operations Manual", level=1)
    for index in range(30):
        document.add_paragraph(
            f"Routine operational paragraph number {index} about daily procedures."
        )
    document.add_heading("Power Requirements", level=1)
    document.add_paragraph(
        "The flux capacitor requires exactly 1.21 gigawatts of power "
        "supplied by the plutonium reactor."
    )
    document.add_heading("Billing", level=1)
    document.add_paragraph(
        "Invoice 0231 was issued for the October plutonium delivery."
    )
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


async def test_docx_through_real_pipeline_then_hybrid_retrieval(
    session: AsyncSession,
    qdrant_collection: None,
) -> None:
    context, workspace = await seed_workspace(session, "e2e-1")
    document = await create_from_upload(
        session,
        context,
        workspace.id,
        filename="manual.docx",
        mime=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        data=fixture_docx(),
    )

    await run_parse(document.id, 1)
    await run_chunk(document.id, 1)
    await run_embed_upsert(document.id, 1)

    await session.refresh(document)
    assert document.status == "indexed"
    assert (document.page_count or 0) >= 1
    workspace_id = workspace.id
    document_id = document.id
    stages = {
        job.stage
        for job in (
            await session.execute(
                select(IngestJob).where(IngestJob.document_id == document.id)
            )
        ).scalars()
    }
    assert stages == {"parse", "chunk", "embed", "upsert"}

    keyword_result = await retrieve(
        session,
        context,
        workspace_id,
        "invoice 0231",
        top_k=5,
    )
    assert any(
        "Invoice 0231" in chunk.text
        for chunk in keyword_result.chunks[:5]
    )

    overlap_result = await retrieve(
        session,
        context,
        workspace_id,
        "gigawatts flux capacitor power",
        top_k=5,
    )
    hit = next(
        chunk
        for chunk in overlap_result.chunks[:5]
        if "1.21 gigawatts" in chunk.text
    )
    assert hit.document_id == document_id
    assert hit.page >= 1
    assert hit.chunk_index >= 0


async def test_empty_and_unsupported_inputs_fail_cleanly(
    session: AsyncSession,
    stack_env: None,
) -> None:
    context, workspace = await seed_workspace(session, "e2e-2")
    empty = await create_from_upload(
        session,
        context,
        workspace.id,
        filename="empty.txt",
        mime="text/plain",
        data=b"",
    )

    with pytest.raises(IngestFailure, match="empty"):
        await run_parse(empty.id, 1)

    await session.refresh(empty)
    assert empty.status == "failed"
    assert "empty" in (empty.error or "")

    unsupported = await create_from_upload(
        session,
        context,
        workspace.id,
        filename="blob.xyz",
        mime="application/octet-stream",
        data=b"\x00\x01",
    )
    with pytest.raises(IngestFailure):
        await run_parse(unsupported.id, 1)

    await session.refresh(unsupported)
    assert unsupported.status == "failed"
    assert unsupported.error
